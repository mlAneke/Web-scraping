import json
from docx import Document
from datetime import datetime as dt
from docx.shared import Pt
today = dt.today()

# print(f'Number of jobs = [filtered.json lines/17]: {jobs}')
location_and_job_number = ""
subject = "Open Jobs at https://www.glassdoor.com/Job/ - " + today.strftime("%B %d, %Y")
keyword_list = f"with following keywords: Test Automation Engineer, Espoo"

# Load the JSON file
with open("job_glassdoor_data.json", "r") as file:
    data = json.load(file)

# Create a new Document
doc = Document()

# Add the document title
doc.add_heading(subject, level=1)
doc.add_paragraph(keyword_list)

# Iterate over the job data
for job in data:
    # Extract the job details
    title = job["title"]
    link = job["link"]
    company = job["company"]
    location = job["location"]

    display_text = "Job Website"

    # Add the job details to the document
    doc.add_heading("Company: " + company)
    
    #doc.add_paragraph("Title: " + title)
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(12)
    doc.add_paragraph("Link: " + link)
    doc.add_paragraph("Location: " + location)
    doc.add_paragraph("----------------------------------")

# Save the document as a DOCX file
doc.save("job_glassdoor_data.docx")
